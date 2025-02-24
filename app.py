from flask import Flask, request, send_file
import os
import uuid
import docx
from docx.shared import Pt
from PyPDF2 import PdfMerger
import pdfkit

app = Flask(__name__)

@app.route("/")
def home():
    return "Le serveur fonctionne ! 🚀"

TEMPLATES_FOLDER = "templates/"
OUTPUT_FOLDER = "output/"
pdfkit_config = pdfkit.configuration(wkhtmltopdf="/usr/bin/wkhtmltopdf")  # Assurez-vous que wkhtmltopdf est installé

# Fonction pour remplacer les variables dans un fichier Word
def replace_variables(template_path, variables):
    doc = docx.Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                paragraph.runs[0].font.size = Pt(11)  # Ajustement de la police si besoin
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in variables.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    temp_filename = os.path.join(OUTPUT_FOLDER, f"temp_{uuid.uuid4()}.docx")
    doc.save(temp_filename)
    return temp_filename

# Fonction pour convertir Word en PDF
def convert_to_pdf(word_path):
    pdf_path = word_path.replace(".docx", ".pdf")
    pdfkit.from_file(word_path, pdf_path, configuration=pdfkit_config)
    return pdf_path

# Fonction pour fusionner plusieurs PDFs
def merge_pdfs(pdf_files, output_path):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    merger.write(output_path)
    merger.close()
    return output_path

@app.route("/generate", methods=["POST"])
def generate_documents():
    template_name = request.form["template"]
    fixed_data = {
        "{ADRESSE}": request.form["adresse"],
        "{NOM DU PROPRIETAIRE}": request.form["nom_proprietaire"],
        "{DESCRIPTION DU BIEN}": request.form["description_bien"],
        "{NOM AGENT}": request.form["nom_agent"],
        "{MAIL}": request.form["mail"],
        "{TELEPHONE}": request.form["telephone"],
    }
    names = request.form["noms"].split(",")  # Liste des noms séparés par une virgule
    
    word_files = []
    pdf_files = []
    
    for name in names:
        variables = fixed_data.copy()
        variables["{NOM}"] = name.strip()
        
        template_name = template_name.lower().replace(" ", "_").replace("é", "e").replace("è", "e").replace("ê", "e")
	word_file = replace_variables(os.path.join(TEMPLATES_FOLDER, template_name), variables)

        word_files.append(word_file)
        
        pdf_file = convert_to_pdf(word_file)
        pdf_files.append(pdf_file)
    
    merged_pdf = os.path.join(OUTPUT_FOLDER, f"final_{uuid.uuid4()}.pdf")
    merge_pdfs(pdf_files, merged_pdf)
    
    return send_file(merged_pdf, as_attachment=True)

if __name__ == "__main__":
    os.makedirs(TEMPLATES_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    app.run(debug=True, host="0.0.0.0", port=5000)
